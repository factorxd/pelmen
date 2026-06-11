# ui/mass_generate_dialog.py
import os
import json
import tempfile
import zipfile
import pandas as pd
from docx.enum.text import WD_BREAK
from PySide6.QtWidgets import (
    QDialog, QVBoxLayout, QHBoxLayout, QPushButton, QFileDialog,
    QTableWidget, QTableWidgetItem, QHeaderView, QComboBox,
    QLabel, QProgressBar, QMessageBox, QGroupBox, QFormLayout,
    QRadioButton, QButtonGroup, QLineEdit, QCheckBox
)
from PySide6.QtCore import Qt

class MassGenerateDialog(QDialog):
    def __init__(self, template, parent=None):
        super().__init__(parent)
        self.template = template
        self.setWindowTitle(f"Массовая генерация: {template.name}")
        self.setMinimumSize(800, 600)
        self.df = None
        self.columns = []
        self.field_mapping = {}
        self.init_ui()

    def init_ui(self):
        layout = QVBoxLayout(self)

        # Группа загрузки файла
        file_group = QGroupBox("1. Выберите файл (CSV или Excel)")
        file_layout = QHBoxLayout(file_group)
        self.file_label = QLabel("Файл не выбран")
        self.select_btn = QPushButton("Выбрать файл")
        file_layout.addWidget(self.file_label)
        file_layout.addWidget(self.select_btn)
        layout.addWidget(file_group)

        # Таблица предпросмотра
        self.table = QTableWidget()
        self.table.setVisible(False)
        layout.addWidget(self.table)

        # Группа сопоставления колонок
        self.mapping_group = QGroupBox("2. Сопоставьте колонки с полями шаблона")
        self.mapping_group.setVisible(False)
        mapping_layout = QFormLayout(self.mapping_group)
        self.mapping_widgets = []
        layout.addWidget(self.mapping_group)

        # Группа настроек генерации
        self.settings_group = QGroupBox("3. Настройки генерации")
        self.settings_group.setVisible(False)
        settings_layout = QVBoxLayout(self.settings_group)

        # Режим генерации
        mode_layout = QHBoxLayout()
        self.single_file_radio = QRadioButton("Один файл (с разрывами страниц)")
        self.multi_file_radio = QRadioButton("Несколько файлов (ZIP)")
        self.multi_file_radio.setChecked(True)
        mode_layout.addWidget(self.single_file_radio)
        mode_layout.addWidget(self.multi_file_radio)
        settings_layout.addLayout(mode_layout)

        # Маска имени файла
        mask_layout = QHBoxLayout()
        mask_layout.addWidget(QLabel("Шаблон имени файла:"))
        self.name_mask_edit = QLineEdit()
        self.name_mask_edit.setPlaceholderText("например: Счет_{doc_number}_{client_name}")
        mask_layout.addWidget(self.name_mask_edit)
        settings_layout.addLayout(mask_layout)

        # Опция ZIP (для нескольких файлов)
        self.zip_check = QCheckBox("Создать ZIP-архив (иначе сохранить в папку)")
        self.zip_check.setChecked(True)
        settings_layout.addWidget(self.zip_check)

        layout.addWidget(self.settings_group)

        # Прогресс
        self.progress = QProgressBar()
        self.progress.setVisible(False)
        layout.addWidget(self.progress)

        # Кнопки
        btn_layout = QHBoxLayout()
        self.generate_btn = QPushButton("Сгенерировать")
        self.generate_btn.setEnabled(False)
        self.cancel_btn = QPushButton("Отмена")
        btn_layout.addStretch()
        btn_layout.addWidget(self.generate_btn)
        btn_layout.addWidget(self.cancel_btn)
        layout.addLayout(btn_layout)

        # Сигналы
        self.select_btn.clicked.connect(self.load_file)
        self.generate_btn.clicked.connect(self.generate)
        self.cancel_btn.clicked.connect(self.reject)

    def load_file(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Выберите файл", "",
            "CSV files (*.csv);;Excel files (*.xlsx *.xls)"
        )
        if not file_path:
            return
        self.file_label.setText(file_path)
        try:
            if file_path.endswith('.csv'):
                self.df = pd.read_csv(file_path, encoding='utf-8')
            else:
                self.df = pd.read_excel(file_path)
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось прочитать файл:\n{str(e)}")
            return
        if self.df.empty:
            QMessageBox.critical(self, "Ошибка", "Файл пуст")
            return
        self.columns = list(self.df.columns)
        self.show_preview()
        self.setup_mapping()
        self.table.setVisible(True)
        self.mapping_group.setVisible(True)
        self.settings_group.setVisible(True)
        self.generate_btn.setEnabled(True)

    def show_preview(self):
        self.table.clear()
        preview = self.df.head(5)
        self.table.setRowCount(len(preview))
        self.table.setColumnCount(len(self.columns))
        self.table.setHorizontalHeaderLabels(self.columns)
        for i, row in preview.iterrows():
            for j, col in enumerate(self.columns):
                item = QTableWidgetItem(str(row[col]))
                self.table.setItem(i, j, item)
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.table.setMaximumHeight(200)

    def setup_mapping(self):
        all_fields = []
        for field in self.template.fields:
            all_fields.append(("field", field.name))
        # для простоты не добавляем поля блоков, так как массовая генерация обычно для простых полей
        # но можно добавить, если нужно
        layout = self.mapping_group.layout()
        for i in reversed(range(layout.count())):
            widget = layout.itemAt(i).widget()
            if widget:
                widget.deleteLater()
        self.mapping_widgets.clear()
        for col in self.columns:
            label = QLabel(f"Колонка '{col}' →")
            combo = QComboBox()
            combo.addItem("(не использовать)", None)
            for typ, field_name in all_fields:
                combo.addItem(field_name, (typ, field_name))
            layout.addRow(label, combo)
            self.mapping_widgets.append((col, combo))

    def generate(self):
        from logic.doc_generator import generate_docx

        if self.df is None:
            return
        mapping = {}
        for col, combo in self.mapping_widgets:
            data = combo.currentData()
            if data is not None:
                typ, field_name = data
                mapping[col] = (typ, field_name)
        if not mapping:
            QMessageBox.warning(self, "Ошибка", "Не выбрано ни одного сопоставления колонок")
            return

        total_rows = len(self.df)
        self.progress.setVisible(True)
        self.progress.setMaximum(total_rows)
        self.generate_btn.setEnabled(False)

        # Определяем режим генерации
        single_file = self.single_file_radio.isChecked()
        name_mask = self.name_mask_edit.text().strip()
        create_zip = self.zip_check.isChecked() if not single_file else False

        if single_file:
            from docx import Document
            from docxcompose.composer import Composer
            import shutil

            doc_list = []
            for idx, row in self.df.iterrows():
                data_dict = {}
                for col, (typ, field_name) in mapping.items():
                    value = row[col]
                    if pd.isna(value):
                        value = ""
                    # Получаем информацию о поле (тип, формат, суффикс)
                    field_info = self.get_field_info(field_name)
                    # Применяем форматирование (дата, число)
                    data_dict[field_name] = self.format_value(value, field_info)
                # Блоки пока игнорируем (если нужны — доработать)
                for block in self.template.blocks:
                    if block.name not in data_dict:
                        data_dict[block.name] = []
                # Генерируем временный DOCX
                with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp2:
                    part_path = tmp2.name
                try:
                    from logic.doc_generator import generate_docx
                    generate_docx(self.template.file_path, data_dict, part_path)
                except Exception as e:
                    QMessageBox.critical(self, "Ошибка", f"Ошибка в строке {idx + 1}:\n{str(e)}")
                    return
                # Добавляем разрыв страницы в конец каждого документа (кроме последнего, добавим позже)
                # Сделаем это после объединения, но проще добавить разрыв в сам временный файл
                if idx < total_rows - 1:  # не последний
                    doc = Document(part_path)
                    doc.add_page_break()
                    doc.save(part_path)
                doc_list.append(part_path)
                self.progress.setValue(idx + 1)

            if doc_list:
                # Объединяем документы
                first_doc = Document(doc_list[0])
                composer = Composer(first_doc)
                for path in doc_list[1:]:
                    next_doc = Document(path)
                    composer.append(next_doc)
                # Сохраняем результат
                final_path = doc_list[0]
                composer.save(final_path)

                save_path, _ = QFileDialog.getSaveFileName(
                    self, "Сохранить документ",
                    f"{self.template.name}_merged.docx",
                    "Word files (*.docx)"
                )
                if save_path:
                    shutil.move(final_path, save_path)
                    QMessageBox.information(self, "Успех", f"Документ сохранён:\n{save_path}")
                else:
                    for p in doc_list:
                        os.unlink(p)
        else:
            # Несколько файлов
            with tempfile.TemporaryDirectory() as tmpdir:
                output_files = []
                for idx, row in self.df.iterrows():
                    data_dict = {}
                    for col, (typ, field_name) in mapping.items():
                        value = row[col]
                        if pd.isna(value):
                            value = ""
                        data_dict[field_name] = str(value)
                    for block in self.template.blocks:
                        if block.name not in data_dict:
                            data_dict[block.name] = []
                    # Генерируем имя файла по маске
                    if name_mask:
                        try:
                            # Простая замена {field} на значение
                            fname = name_mask
                            for field_name in data_dict:
                                fname = fname.replace(f"{{{field_name}}}", str(data_dict[field_name]))
                            # Удаляем недопустимые символы
                            fname = "".join(c for c in fname if c.isalnum() or c in "._- ")
                        except:
                            fname = f"doc_{idx+1}"
                    else:
                        fname = f"doc_{idx+1}"
                    out_path = os.path.join(tmpdir, f"{fname}.docx")
                    try:
                        generate_docx(self.template.file_path, data_dict, out_path)
                        output_files.append(out_path)
                    except Exception as e:
                        QMessageBox.critical(self, "Ошибка", f"Ошибка в строке {idx+1}:\n{str(e)}")
                        self.progress.setVisible(False)
                        self.generate_btn.setEnabled(True)
                        return
                    self.progress.setValue(idx+1)
                if create_zip:
                    zip_path, _ = QFileDialog.getSaveFileName(self, "Сохранить архив", f"{self.template.name}_mass.zip", "ZIP files (*.zip)")
                    if zip_path:
                        with zipfile.ZipFile(zip_path, 'w') as zipf:
                            for fpath in output_files:
                                zipf.write(fpath, os.path.basename(fpath))
                        QMessageBox.information(self, "Успех", f"Создано {len(output_files)} документов\nСохранено в {zip_path}")
                    else:
                        QMessageBox.warning(self, "Отменено", "Архив не сохранён")
                else:
                    folder = QFileDialog.getExistingDirectory(self, "Выберите папку для сохранения")
                    if folder:
                        import shutil
                        for fpath in output_files:
                            shutil.move(fpath, os.path.join(folder, os.path.basename(fpath)))
                        QMessageBox.information(self, "Успех", f"Сохранено {len(output_files)} документов в папку {folder}")
                    else:
                        QMessageBox.warning(self, "Отменено", "Сохранение отменено")
        self.accept()

    def get_field_info(self, field_name):
        """Возвращает информацию о поле (тип, формат и т.д.) из настроек шаблона"""
        key = f"field:{field_name}"
        data = self.parent().display_names.get(self.template.id, {}).get(key, {})
        if isinstance(data, str):
            return {"type": "text", "format": ""}
        return {
            "type": data.get("type", "text"),
            "format": data.get("format", "")
        }

    def format_value(self, value, field_info):
        """Форматирует значение в соответствии с типом поля и форматом"""
        field_type = field_info.get("type", "text")
        fmt = field_info.get("format", "")

        if pd.isna(value):
            return ""

        if field_type == "date":
            try:
                dt = pd.to_datetime(value)
                if fmt:
                    # Преобразуем шаблон формата, например "DD.MM.YYYY" в "%d.%m.%Y"
                    # Но пользователь может хранить формат уже в виде "%d.%m.%Y"
                    # Поэтому попробуем оба варианта
                    if '%' in fmt:
                        # Уже strftime-формат
                        return dt.strftime(fmt)
                    else:
                        # Конвертируем из нашего формата в strftime
                        fmt = fmt.replace("DD", "%d").replace("MM", "%m").replace("YYYY", "%Y")
                        return dt.strftime(fmt)
                else:
                    return dt.strftime("%d.%m.%Y")
            except:
                return str(value)
        elif field_type == "number":
            try:
                num = float(value)
                if fmt:
                    try:
                        formatted = fmt.format(num)
                    except:
                        formatted = str(num)
                else:
                    formatted = f"{num:.2f}" if isinstance(num, float) and num % 1 != 0 else str(int(num))

                return formatted
            except:
                return str(value)
        elif field_type == "bool":
            # Преобразуем в булево
            if isinstance(value, bool):
                return value
            if isinstance(value, str):
                return value.lower() in ("true", "1", "yes", "да")
            return bool(value)
        else:
            return str(value)